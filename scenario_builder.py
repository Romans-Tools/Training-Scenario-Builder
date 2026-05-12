#!/usr/bin/env python3
"""Local-only Training Scenario Builder CLI.

- Runs fully offline.
- Keeps all working data in memory only.
- Exports only when explicitly requested.
"""

from __future__ import annotations

import random
from dataclasses import dataclass
from datetime import datetime, timedelta
from textwrap import fill
from typing import Dict, List


VARIABLE_POOLS: Dict[str, List[str]] = {
    "operational_periods": ["morning", "afternoon", "evening", "overnight"],
    "weather_states": ["VFR with increasing winds", "IFR ceilings", "thunderstorm watch", "extreme heat advisory"],
    "media_pressure": ["low", "moderate", "high", "intense"],
    "comms_status": ["stable", "degraded repeater coverage", "intermittent data loss", "primary channel outage"],
    "partner_cadence": ["hourly county EOC calls", "30-minute unified command sync", "ad hoc agency updates", "joint information center rhythm"],
}


TEMPLATES: Dict[str, Dict[str, List[str]]] = {
    "mission base operations": {
        "objectives": [
            "Establish Civil Air Patrol incident support workflows using ICS structure.",
            "Synchronize aircrew, ground team, and mission staff communications.",
            "Prioritize safety, accountability, and sortie readiness under time pressure.",
        ],
        "injects": [
            "Weather shifts force rapid sortie replanning and crew brief updates.",
            "An overdue aircraft report requires immediate mission base coordination.",
            "Fuel and vehicle availability constraints impact launch timelines.",
        ],
    },
    "search and rescue": {
        "objectives": [
            "Coordinate CAP air and ground assets for time-critical search coverage.",
            "Use standardized mission documentation and resource tracking.",
            "Integrate partner-agency updates into operational decisions.",
        ],
        "injects": [
            "ELT signal quality fluctuates and creates conflicting search vectors.",
            "A ground team requests medevac support from a remote trailhead.",
            "Airspace restrictions suddenly reduce available search sectors.",
        ],
    },
    "disaster relief": {
        "objectives": [
            "Plan CAP support for post-disaster assessment and logistics coordination.",
            "Manage tasking priorities from emergency management partners.",
            "Maintain mission continuity while protecting volunteer endurance.",
        ],
        "injects": [
            "A storm-damaged comm relay degrades mission base connectivity.",
            "A shelter location needs urgent aerial imagery verification.",
            "Competing task requests exceed available qualified crews.",
        ],
    },
    "public affairs": {
        "objectives": [
            "Deliver timely, accurate CAP mission updates to the public and media.",
            "Protect operational security while maintaining transparency and trust.",
            "Coordinate unified messaging with incident command and partner agencies.",
        ],
        "injects": [
            "A viral post misidentifies CAP activity as unauthorized flight operations.",
            "A local station requests a live briefing before facts are fully confirmed.",
            "Family members post conflicting details about missing-person status online.",
        ],
    },
    "cadet activity support": {
        "objectives": [
            "Apply CAP safety and supervision standards during high-visibility events.",
            "Coordinate logistics and communications across senior and cadet staff.",
            "Respond to emerging issues while preserving training objectives.",
        ],
        "injects": [
            "Transportation delays threaten the event start timeline.",
            "A heat advisory requires immediate risk-control adjustments.",
            "Parents and community members request simultaneous status updates.",
        ],
    },
}



@dataclass
class Inputs:
    training_type: str
    audience: str
    duration_minutes: int
    difficulty: str
    objectives: List[str]
    setting: str
    participants: int
    constraints: List[str]


def ask(prompt: str, default: str | None = None) -> str:
    suffix = f" [{default}]" if default else ""
    value = input(f"{prompt}{suffix}: ").strip()
    return value if value else (default or "")


def parse_list(raw: str) -> List[str]:
    return [item.strip() for item in raw.split(",") if item.strip()]


def collect_inputs() -> Inputs:
    print("\n=== Public Information Officer Builder (Local-Only CLI) ===")
    print("All data remains in-memory and is not saved unless you export.\n")

    training_type = ask(
        "Training type (mission base operations/search and rescue/disaster relief/public affairs/cadet activity support)",
        "mission base operations",
    ).lower()
    if training_type not in TEMPLATES:
        print("Unknown type; using 'mission base operations' template defaults.")
        training_type = "mission base operations"

    audience = ask("Audience", "CAP PIOs and mission staff")
    duration_minutes = int(ask("Duration in minutes", "120"))
    difficulty = ask("Difficulty (beginner/intermediate/advanced)", "intermediate")
    objectives = parse_list(ask("Objectives (comma-separated, optional)", ""))
    setting = ask("Setting", "CAP mission base")
    participants = int(ask("Number of participants", "12"))
    constraints = parse_list(
        ask("Constraints (comma-separated, optional)", "limited staffing, media pressure")
    )

    return Inputs(
        training_type=training_type,
        audience=audience,
        duration_minutes=duration_minutes,
        difficulty=difficulty,
        objectives=objectives,
        setting=setting,
        participants=participants,
        constraints=constraints,
    )


def build_scenario(data: Inputs) -> Dict[str, object]:
    template = TEMPLATES[data.training_type]
    objectives = data.objectives or random.sample(
        template["objectives"], k=min(3, len(template["objectives"]))
    )

    title = (
        f"Public Information Officer Scenario: {data.difficulty.title()} "
        f"{data.setting} {data.training_type.title()} Exercise"
    )

    num_roles = min(max(data.participants, 4), 10)
    roles = [
        f"Role {i+1}: {'Public Information Officer' if i == 0 else 'Mission Staff Participant'}"
        for i in range(num_roles)
    ]

    start = datetime(2000, 1, 1, 9, 0)
    step = max(10, data.duration_minutes // 5)
    timeline = [
        f"{(start + timedelta(minutes=i*step)).strftime('%H:%M')} - Phase {i+1}"
        for i in range(6)
    ]

    scenario_variables = [
        f"Operational period: {random.choice(VARIABLE_POOLS['operational_periods'])}.",
        f"Weather factor: {random.choice(VARIABLE_POOLS['weather_states'])}.",
        f"Media pressure: {random.choice(VARIABLE_POOLS['media_pressure'])}.",
        f"Communications status: {random.choice(VARIABLE_POOLS['comms_status'])}.",
        f"Partner coordination cadence: {random.choice(VARIABLE_POOLS['partner_cadence'])}.",
    ]

    injects = random.sample(template["injects"], k=min(3, len(template["injects"]))) + [
        f"Constraint pressure: {c}." for c in data.constraints[:2]
    ]

    expected_actions = [
        "Establish clear command/decision roles within first phase.",
        "Produce a shared action plan with owner and deadline for each task.",
        "Communicate status updates at defined intervals.",
        "Document assumptions, risks, and mitigation actions.",
        "Update the public narrative whenever operational facts materially change.",
    ]

    facilitator_notes = [
        "Escalate complexity only after participants demonstrate baseline coordination.",
        "Observe decision rationale, not only outcomes.",
        "Capture notable communication failures and recoveries for debrief.",
    ]

    rubric = [
        "Command & Coordination (1-5): role clarity, alignment, control.",
        "Communication Quality (1-5): timeliness, accuracy, consistency.",
        "Decision-Making (1-5): risk balance, prioritization, adaptability.",
        "Execution (1-5): task completion, accountability, follow-through.",
    ]

    aar = [
        "What signals were recognized early, late, or missed?",
        "Which decisions had the greatest downstream impact?",
        "Where did communication break down and why?",
        "What plan, policy, or training updates should be made?",
        "What should be sustained, improved, and stopped?",
    ]

    return {
        "title": title,
        "overview": (
            f"A {data.duration_minutes}-minute Civil Air Patrol {data.training_type} exercise for "
            f"{data.audience} in a {data.setting} setting at {data.difficulty} difficulty."
        ),
        "training_objectives": objectives,
        "participant_roles": roles,
        "timeline": timeline,
        "scenario_variables": scenario_variables,
        "injects_events": injects,
        "expected_actions": expected_actions,
        "facilitator_notes": facilitator_notes,
        "evaluation_rubric": rubric,
        "after_action_review_questions": aar,
        "constraints": data.constraints,
    }


def render_markdown(pkg: Dict[str, object]) -> str:
    def md_list(items: List[str]) -> str:
        return "\n".join(f"- {item}" for item in items)

    return f"""# {pkg['title']}

## Overview
{pkg['overview']}

## Training Objectives
{md_list(pkg['training_objectives'])}

## Participant Roles
{md_list(pkg['participant_roles'])}

## Timeline
{md_list(pkg['timeline'])}

## Scenario Variables
{md_list(pkg['scenario_variables'])}

## Injects/Events
{md_list(pkg['injects_events'])}

## Expected Participant Actions
{md_list(pkg['expected_actions'])}

## Facilitator Notes
{md_list(pkg['facilitator_notes'])}

## Evaluation Rubric
{md_list(pkg['evaluation_rubric'])}

## After-Action Review Questions
{md_list(pkg['after_action_review_questions'])}
"""


def preview(pkg: Dict[str, object]) -> None:
    print("\n=== Preview ===")
    print(pkg["title"])
    print(fill(pkg["overview"], width=88))
    print("\nObjectives:")
    for o in pkg["training_objectives"][:3]:
        print(f"  • {o}")
    print("\nInjects:")
    for ev in pkg["injects_events"][:3]:
        print(f"  • {ev}")


def export_markdown(content: str, out_path: str) -> None:
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)


def export_docx(content: str, out_path: str) -> None:
    from docx import Document

    doc = Document()
    for line in content.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(out_path)


def export_pdf(content: str, out_path: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(out_path, pagesize=letter)
    width, height = letter
    y = height - 50
    for raw_line in content.splitlines():
        lines = [raw_line[i : i + 95] for i in range(0, len(raw_line), 95)] or [""]
        for line in lines:
            c.drawString(40, y, line)
            y -= 14
            if y < 40:
                c.showPage()
                y = height - 50
    c.save()


def main() -> None:
    scenario_inputs = collect_inputs()
    scenario = build_scenario(scenario_inputs)
    markdown = render_markdown(scenario)

    preview(scenario)

    if ask("\nExport scenario? (y/n)", "y").lower() != "y":
        print("No export selected. Data cleared from memory on exit.")
        return

    fmt = ask("Format (md/pdf/docx)", "md").lower()
    out = ask("Output file path", f"scenario_export.{fmt}")

    if fmt == "md":
        export_markdown(markdown, out)
    elif fmt == "docx":
        try:
            export_docx(markdown, out)
        except ModuleNotFoundError:
            print("Missing dependency: python-docx. Install locally to enable DOCX export.")
            return
    elif fmt == "pdf":
        try:
            export_pdf(markdown, out)
        except ModuleNotFoundError:
            print("Missing dependency: reportlab. Install locally to enable PDF export.")
            return
    else:
        print("Unsupported format. Use md, pdf, or docx.")
        return

    print(f"Export complete: {out}")
    print("Scenario data will be cleared from memory when the program exits.")


if __name__ == "__main__":
    main()
